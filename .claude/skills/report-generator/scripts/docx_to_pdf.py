#!/usr/bin/env python3
"""
DOCX to PDF converter with professional xOS branding.
Uses python-docx to read and reportlab to generate.
"""

import sys, os
from pathlib import Path
from xml.sax.saxutils import escape as xml_escape

from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch, mm
from reportlab.lib.colors import HexColor, black, white
from reportlab.platypus import (
    BaseDocTemplate, PageTemplate, Frame, Paragraph, Spacer,
    PageBreak, Table as RLTable, TableStyle, KeepTogether, Flowable,
    NextPageTemplate
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# ── Brand (Bjarn Brunenberg Brand Colours) ──
BRAND_DARK = "#022147"       # Deep Navy — darkest backgrounds, gradient edges
BRAND_BLUE = "#135BB4"       # Brand Blue — hero backgrounds, primary brand colour
BRAND_HIGHLIGHT = "#4DE2FE"  # Electric Cyan — accent ONLY, CTAs, highlights (max 10-15%)
BRAND_SILVER = "#D9D9D9"     # Silver — logo, name typography on dark backgrounds
BRAND_LIGHT = "#EBF4FF"      # Light blue tint for alternating rows / light backgrounds
BRAND_WHITE = "#FFFFFF"      # White — headlines, body text on dark backgrounds
BRAND_GRAY = "#022147"       # Body text colour = Deep Navy for readability on white

# ── Fonts (Poppins — closest to DM Sans available) ──
pdfmetrics.registerFont(TTFont("Lib", "/usr/share/fonts/truetype/google-fonts/Poppins-Regular.ttf"))
pdfmetrics.registerFont(TTFont("LibB", "/usr/share/fonts/truetype/google-fonts/Poppins-Bold.ttf"))
pdfmetrics.registerFont(TTFont("LibI", "/usr/share/fonts/truetype/google-fonts/Poppins-Italic.ttf"))
pdfmetrics.registerFont(TTFont("LibBI", "/usr/share/fonts/truetype/google-fonts/Poppins-BoldItalic.ttf"))
pdfmetrics.registerFontFamily('Lib', normal='Lib', bold='LibB', italic='LibI', boldItalic='LibBI')
# Medium weight for subheadings
pdfmetrics.registerFont(TTFont("LibM", "/usr/share/fonts/truetype/google-fonts/Poppins-Medium.ttf"))
pdfmetrics.registerFont(TTFont("LibMI", "/usr/share/fonts/truetype/google-fonts/Poppins-MediumItalic.ttf"))

PAGE_W, PAGE_H = A4
MARGIN = 0.75 * inch
CONTENT_W = PAGE_W - 2 * MARGIN

# ── Styles ──
def mk_style(name, **kw):
    defaults = dict(fontName='Lib', fontSize=10, textColor=HexColor(BRAND_GRAY), leading=13, spaceAfter=6)
    defaults.update(kw)
    return ParagraphStyle(name, **defaults)

S_BODY = mk_style('Body', fontSize=10.5, leading=14, spaceAfter=8)
S_H1 = mk_style('H1', fontName='LibB', fontSize=22, textColor=HexColor(BRAND_DARK), spaceBefore=6, spaceAfter=10, leading=26)
S_H2 = mk_style('H2', fontName='LibB', fontSize=15, textColor=HexColor(BRAND_BLUE), spaceBefore=10, spaceAfter=8, leading=19)
S_H3 = mk_style('H3', fontName='LibM', fontSize=12, textColor=HexColor(BRAND_DARK), spaceBefore=8, spaceAfter=6, leading=16)
S_BULLET = mk_style('Bullet', fontSize=10.5, leading=14, spaceAfter=4, leftIndent=24, bulletIndent=12)
S_TCELL = mk_style('TCell', fontSize=9, leading=12, spaceAfter=0, spaceBefore=0)
S_TCELL_B = mk_style('TCellB', fontName='LibB', fontSize=9, leading=12, spaceAfter=0, spaceBefore=0)
S_TCELL_W = mk_style('TCellW', fontName='LibB', fontSize=9, leading=12, textColor=white, spaceAfter=0, spaceBefore=0)
S_COVER_TITLE = mk_style('CTitle', fontName='LibB', fontSize=42, textColor=white, alignment=TA_CENTER, leading=48, spaceAfter=8)
S_COVER_SUB = mk_style('CSub', fontName='Lib', fontSize=18, textColor=HexColor(BRAND_SILVER), alignment=TA_CENTER, leading=24, spaceAfter=6)
S_COVER_TAG = mk_style('CTag', fontSize=10, textColor=HexColor(BRAND_SILVER), alignment=TA_CENTER, leading=15, spaceAfter=6)
S_COVER_META = mk_style('CMeta', fontName='LibM', fontSize=11, textColor=HexColor(BRAND_SILVER), alignment=TA_CENTER, leading=14, spaceAfter=4)
S_COVER_CONF = mk_style('CConf', fontName='LibB', fontSize=11, textColor=HexColor(BRAND_HIGHLIGHT), alignment=TA_CENTER, leading=14, spaceAfter=4)


def esc(text):
    """Escape text for reportlab XML paragraphs."""
    if not text:
        return ""
    return xml_escape(str(text), {'"': '&quot;'})


def runs_to_markup(para):
    """Convert paragraph runs to reportlab XML markup preserving bold/italic."""
    parts = []
    for run in para.runs:
        t = esc(run.text)
        if not t:
            continue
        if run.bold and run.italic:
            parts.append(f'<b><i>{t}</i></b>')
        elif run.bold:
            parts.append(f'<b>{t}</b>')
        elif run.italic:
            parts.append(f'<i>{t}</i>')
        else:
            parts.append(t)
    return ''.join(parts) if parts else esc(para.text)


def cell_bg(cell):
    """Extract cell background color hex string or None."""
    try:
        tcPr = cell._tc.tcPr
        if tcPr is None:
            return None
        ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        shd = tcPr.find(f'{ns}shd')
        if shd is None:
            return None
        fill = shd.get(f'{ns}fill', '')
        if not fill or fill.upper() == 'AUTO':
            return None
        return fill.lower()
    except:
        return None


def map_color(hex_color):
    """Map a docx hex color to our brand palette."""
    if not hex_color:
        return None
    h = hex_color.lower().replace('#', '')
    # Dark backgrounds → Deep Navy
    if h in ('1a1a2e', '2a2a4e', '1a1a3e', '000000', '16213e', '022147'):
        return BRAND_DARK
    # Red/pink — keep for semantic use (Kill, Critical, High risk)
    if h in ('fee2e2', 'ffcccc', 'ff9999', 'ffc7ce'):
        return "#FEE2E2"
    # Yellow/amber — keep for semantic use (Pivot, Medium risk)
    if h in ('fef3c7', 'ffe6a6', 'ffcc80', 'fff2cc', 'fce4b6'):
        return "#FEF3C7"
    # Blue
    if h in ('dbeafe', 'b3e5fc', 'cfe2f3', 'd6e4f0', 'dae3f3'):
        return "#DBEAFE"
    # Green — keep for semantic use (Go, Low risk)
    if h in ('d1fae5', 'a8e6cf', 'd9ead3', 'e2efda'):
        return "#D1FAE5"
    # Light/white backgrounds - skip
    if h in ('ffffff', 'f5f5fa', 'f2f2f2', 'f5f5f5', 'auto', 'ebf4ff'):
        return None
    # Return the original as a hex
    try:
        int(h, 16)
        return f"#{h}"
    except:
        return None


def is_callout_table(table):
    """Check if a table is a 1x2 callout/accent box (accent bar + content)."""
    if len(table.rows) == 1 and len(table.columns) == 2:
        first_cell_text = table.rows[0].cells[0].text.strip()
        if not first_cell_text:  # Empty accent bar cell
            return True
    return False


def is_heading(para):
    """Check if paragraph is a heading."""
    return para.style and para.style.name and 'Heading' in para.style.name


def heading_level(para):
    if 'Heading 1' in para.style.name or para.style.name == 'Heading1':
        return 1
    elif 'Heading 2' in para.style.name or para.style.name == 'Heading2':
        return 2
    elif 'Heading 3' in para.style.name or para.style.name == 'Heading3':
        return 3
    return 4


def is_list_para(para):
    """Check if paragraph is a bullet/numbered list item."""
    try:
        if para.style and 'List' in para.style.name:
            return True
        pPr = para._element.pPr
        if pPr is not None:
            ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            if pPr.find(f'{ns}numPr') is not None:
                return True
    except:
        pass
    return False


class AccentBox(Flowable):
    """A callout box with left accent border and light background."""

    def __init__(self, text_parts, width=CONTENT_W, accent_color=BRAND_HIGHLIGHT, bg_color=BRAND_LIGHT):
        super().__init__()
        self.text_parts = text_parts  # list of (text, is_bold)
        self.box_width = width
        self.accent_color = accent_color
        self.bg_color = bg_color
        self.padding = 12
        self.accent_width = 4

        # Pre-build paragraphs to calculate height
        self.paragraphs = []
        for txt, is_bold in text_parts:
            style = mk_style('BoxText',
                fontName='LibB' if is_bold else 'Lib',
                fontSize=10, leading=14, spaceAfter=4,
                textColor=HexColor(BRAND_BLUE) if is_bold else HexColor(BRAND_DARK))
            p = Paragraph(esc(txt), style)
            self.paragraphs.append(p)

        # Calculate height
        avail_w = self.box_width - self.accent_width - 3 * self.padding
        total_h = 2 * self.padding
        for p in self.paragraphs:
            w, h = p.wrap(avail_w, 1000)
            total_h += h + 2
        self.box_height = total_h

    def wrap(self, availWidth, availHeight):
        return self.box_width, self.box_height

    def draw(self):
        c = self.canv
        # Background
        c.setFillColor(HexColor(self.bg_color))
        c.roundRect(0, 0, self.box_width, self.box_height, 4, fill=1, stroke=0)
        # Accent bar
        c.setFillColor(HexColor(self.accent_color))
        c.rect(0, 0, self.accent_width, self.box_height, fill=1, stroke=0)
        # Text
        x = self.accent_width + self.padding
        y = self.box_height - self.padding
        avail_w = self.box_width - self.accent_width - 3 * self.padding
        for p in self.paragraphs:
            w, h = p.wrap(avail_w, 1000)
            p.drawOn(c, x, y - h)
            y -= h + 2


def build_table(docx_table, page_width=CONTENT_W):
    """Convert a docx table to a reportlab Table with proper styling."""
    rows = docx_table.rows
    if not rows:
        return None

    ncols = len(rows[0].cells)
    nrows = len(rows)

    # Build data as Paragraph objects
    data = []
    cell_colors = {}  # (row, col) -> mapped color

    for ri, row in enumerate(rows):
        row_data = []
        for ci, cell in enumerate(row.cells):
            # Get cell text with formatting
            cell_text_parts = []
            for p in cell.paragraphs:
                cell_text_parts.append(runs_to_markup(p))
            full_text = '<br/>'.join(cell_text_parts) if cell_text_parts else ''

            # Determine style
            bg = cell_bg(cell)
            mapped = map_color(bg) if bg else None
            if mapped:
                cell_colors[(ri, ci)] = mapped

            is_header = (ri == 0)
            is_dark = (mapped == BRAND_DARK) if mapped else False

            if is_header or is_dark:
                style = S_TCELL_W
            elif any(p.runs and any(r.bold for r in p.runs) for p in cell.paragraphs):
                style = S_TCELL_B
            else:
                style = S_TCELL

            p = Paragraph(full_text, style)
            row_data.append(p)
        data.append(row_data)

    # Calculate column widths
    if ncols <= 2:
        col_widths = [page_width / ncols] * ncols
    elif ncols == 3:
        col_widths = [page_width * 0.35, page_width * 0.30, page_width * 0.35]
    elif ncols == 4:
        col_widths = [page_width * 0.28, page_width * 0.24, page_width * 0.24, page_width * 0.24]
    elif ncols == 5:
        col_widths = [page_width * 0.08, page_width * 0.27, page_width * 0.15, page_width * 0.25, page_width * 0.25]
    elif ncols == 6:
        col_widths = [page_width * 0.06, page_width * 0.22, page_width * 0.12, page_width * 0.12, page_width * 0.24, page_width * 0.24]
    elif ncols == 7:
        col_widths = [page_width * 0.05, page_width * 0.08, page_width * 0.20, page_width * 0.10, page_width * 0.10, page_width * 0.24, page_width * 0.23]
    else:
        col_widths = [page_width / ncols] * ncols

    # Ensure total width matches
    total = sum(col_widths)
    col_widths = [w * page_width / total for w in col_widths]

    tbl = RLTable(data, colWidths=col_widths, repeatRows=1)

    # Base style
    style_cmds = [
        ('FONTNAME', (0, 0), (-1, -1), 'Lib'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('TEXTCOLOR', (0, 0), (-1, -1), HexColor(BRAND_DARK)),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('GRID', (0, 0), (-1, -1), 0.5, HexColor(BRAND_SILVER)),
        # Header row — Deep Navy background
        ('BACKGROUND', (0, 0), (-1, 0), HexColor(BRAND_DARK)),
        ('TEXTCOLOR', (0, 0), (-1, 0), white),
        ('FONTNAME', (0, 0), (-1, 0), 'LibB'),
    ]

    # Alternating row colors — light blue tint
    for ri in range(1, nrows):
        if ri % 2 == 0:
            style_cmds.append(('BACKGROUND', (0, ri), (-1, ri), HexColor(BRAND_LIGHT)))

    # Cell-specific colors
    for (ri, ci), color in cell_colors.items():
        if ri == 0:
            continue  # Header already styled
        style_cmds.append(('BACKGROUND', (ci, ri), (ci, ri), HexColor(color)))
        if color == BRAND_DARK:
            style_cmds.append(('TEXTCOLOR', (ci, ri), (ci, ri), white))

    tbl.setStyle(TableStyle(style_cmds))
    return tbl


def convert(input_path, output_path):
    """Main conversion function."""
    doc = Document(input_path)
    doc_title = Path(input_path).stem.replace('_', ' ').replace('xOS ', 'xOS ')

    # ── Detect cover page paragraphs (before first heading) ──
    cover_end = 0
    for i, p in enumerate(doc.paragraphs):
        if is_heading(p):
            cover_end = i
            break
    else:
        cover_end = len(doc.paragraphs)

    cover_texts = []
    for i in range(cover_end):
        t = doc.paragraphs[i].text.strip()
        if t:
            cover_texts.append(t)

    # ── Build PDF ──
    pdf = BaseDocTemplate(output_path, pagesize=A4,
        leftMargin=MARGIN, rightMargin=MARGIN, topMargin=MARGIN, bottomMargin=MARGIN)

    frame = Frame(MARGIN, MARGIN, CONTENT_W, PAGE_H - 2*MARGIN, id='main')
    cover_frame = Frame(MARGIN, MARGIN, CONTENT_W, PAGE_H - 2*MARGIN, id='cover')

    def cover_bg(canvas, doc):
        canvas.saveState()
        # Solid Deep Navy background
        canvas.setFillColor(HexColor(BRAND_DARK))
        canvas.rect(0, 0, PAGE_W, PAGE_H, fill=1, stroke=0)
        canvas.restoreState()

    def normal_bg(canvas, doc):
        canvas.saveState()
        # White page background (explicit)
        canvas.setFillColor(white)
        canvas.rect(0, 0, PAGE_W, PAGE_H, fill=1, stroke=0)
        # Header bar — Deep Navy
        canvas.setFillColor(HexColor(BRAND_DARK))
        canvas.rect(0, PAGE_H - 30, PAGE_W, 30, fill=1, stroke=0)
        canvas.setFont('LibB', 8)
        canvas.setFillColor(white)
        canvas.drawString(MARGIN, PAGE_H - 22, f"xOS | {doc_title}")
        # Thin cyan accent line below header
        canvas.setStrokeColor(HexColor(BRAND_HIGHLIGHT))
        canvas.setLineWidth(1.5)
        canvas.line(0, PAGE_H - 30, PAGE_W, PAGE_H - 30)
        # Footer
        canvas.setFont('Lib', 8)
        canvas.setFillColor(HexColor(BRAND_DARK))
        canvas.drawString(MARGIN, 24, f"CONFIDENTIAL | Page {doc.page - 1}")
        # Thin footer line
        canvas.setStrokeColor(HexColor(BRAND_SILVER))
        canvas.setLineWidth(0.5)
        canvas.line(MARGIN, 38, PAGE_W - MARGIN, 38)
        canvas.restoreState()

    cover_tmpl = PageTemplate('cover', frames=[cover_frame], onPage=cover_bg)
    normal_tmpl = PageTemplate('normal', frames=[frame], onPage=normal_bg)
    pdf.addPageTemplates([cover_tmpl, normal_tmpl])

    story = []

    # ── Cover page ──
    story.append(Spacer(1, 2.2 * inch))

    if len(cover_texts) >= 1:
        # Check if first text is a label like "EXPERIMENT ROADMAP" or "PRODUCT VISION"
        first = cover_texts[0]
        if first.isupper() or len(first.split()) <= 4:
            # Spaced-out label
            spaced = '     '.join(' '.join(word) for word in first.split())
            story.append(Paragraph(esc(spaced), mk_style('CLabel', fontName='LibM', fontSize=12,
                textColor=HexColor(BRAND_HIGHLIGHT), alignment=TA_CENTER, leading=16, spaceAfter=12, letterSpacing=4)))
        else:
            story.append(Paragraph(esc(first), S_COVER_TITLE))

    if len(cover_texts) >= 2:
        story.append(Paragraph(esc(cover_texts[1]), S_COVER_TITLE))

    if len(cover_texts) >= 3:
        story.append(Paragraph(esc(cover_texts[2]), S_COVER_SUB))

    # Space before taglines
    story.append(Spacer(1, 20))

    # Remaining cover text — split into taglines vs meta
    taglines = []
    meta_lines = []
    for ct in cover_texts[3:]:
        if ct.upper() == 'CONFIDENTIAL' or any(kw in ct.lower() for kw in ['bjarn', 'february', '2026', '2025', 'prepared for', 'github']):
            meta_lines.append(ct)
        else:
            taglines.append(ct)

    # Taglines with good spacing
    for tl in taglines:
        story.append(Paragraph(esc(tl), S_COVER_TAG))

    # Space between taglines and meta info
    if taglines and meta_lines:
        story.append(Spacer(1, 24))

    # Meta lines
    for ml in meta_lines:
        if ml.upper() == 'CONFIDENTIAL':
            story.append(Paragraph(esc(ml), S_COVER_CONF))
        else:
            story.append(Paragraph(esc(ml), S_COVER_META))

    story.append(NextPageTemplate('normal'))
    story.append(PageBreak())

    # ── Main content ──
    # We iterate through body elements to maintain order of paragraphs and tables
    body = doc.element.body
    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    first_heading_seen = False

    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'p':
            para = para_map.get(child)
            if para is None:
                continue

            # Skip cover paragraphs
            idx = None
            for i, dp in enumerate(doc.paragraphs):
                if dp._element is child:
                    idx = i
                    break
            if idx is not None and idx < cover_end:
                continue

            text = para.text.strip()

            if is_heading(para):
                lvl = heading_level(para)
                # Page break before H1 (except the very first one)
                if lvl == 1:
                    if first_heading_seen:
                        story.append(PageBreak())
                    first_heading_seen = True

                style = [None, S_H1, S_H2, S_H3, S_H3][lvl]
                story.append(Paragraph(runs_to_markup(para), style))

            elif is_list_para(para) and text:
                story.append(Paragraph(f"• {runs_to_markup(para)}", S_BULLET))

            elif text:
                story.append(Paragraph(runs_to_markup(para), S_BODY))

            else:
                story.append(Spacer(1, 4))

        elif tag == 'tbl':
            tbl_obj = table_map.get(child)
            if tbl_obj is None:
                continue

            # Check if it's a code block (1x1 table with monospace content)
            if len(tbl_obj.rows) == 1 and len(tbl_obj.columns) == 1 and not is_callout_table(tbl_obj):
                cell = tbl_obj.rows[0].cells[0]
                code_text = cell.text
                lines = code_text.split('\n')
                # Build multi-row table so it can split across pages
                S_CODE = mk_style('Code', fontName='Courier', fontSize=8.5, leading=11,
                    spaceAfter=0, spaceBefore=0, textColor=HexColor(BRAND_DARK))
                code_rows = []
                for line in lines:
                    code_rows.append([Paragraph(esc(line) if line.strip() else '&nbsp;', S_CODE)])
                if code_rows:
                    code_tbl = RLTable(code_rows, colWidths=[CONTENT_W])
                    code_tbl.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, -1), HexColor("#F0F4FA")),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('LEFTPADDING', (0, 0), (-1, -1), 10),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                        ('TOPPADDING', (0, 0), (-1, -1), 2),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
                        ('TOPPADDING', (0, 0), (-1, 0), 8),
                        ('BOTTOMPADDING', (0, -1), (-1, -1), 8),
                        ('LINEABOVE', (0, 0), (-1, 0), 0.5, HexColor(BRAND_SILVER)),
                        ('LINEBELOW', (0, -1), (-1, -1), 0.5, HexColor(BRAND_SILVER)),
                        ('LINEBEFORE', (0, 0), (0, -1), 0.5, HexColor(BRAND_SILVER)),
                        ('LINEAFTER', (-1, 0), (-1, -1), 0.5, HexColor(BRAND_SILVER)),
                    ]))
                    story.append(Spacer(1, 6))
                    story.append(code_tbl)
                    story.append(Spacer(1, 8))
                continue

            # Check if it's a callout box (1x1 or 1x2 table)
            if is_callout_table(tbl_obj):
                # For 1x2 tables, content is in the second cell; for 1x1, it's in the first
                cell = tbl_obj.rows[0].cells[-1]  # Last cell has content
                parts = []
                for p in cell.paragraphs:
                    t = p.text.strip()
                    if t:
                        has_bold = any(r.bold for r in p.runs) if p.runs else False
                        parts.append((t, has_bold))
                if parts:
                    # Build callout as a 2-column table (accent bar + content) for page-split support
                    content_paras = []
                    for txt_part, is_bold in parts:
                        style = mk_style('BoxTxt',
                            fontName='LibB' if is_bold else 'Lib',
                            fontSize=10, leading=14, spaceAfter=4,
                            textColor=HexColor(BRAND_BLUE) if is_bold else HexColor(BRAND_DARK))
                        content_paras.append(Paragraph(esc(txt_part), style))
                    accent_cell_w = 4
                    content_cell_w = CONTENT_W - accent_cell_w
                    callout_tbl = RLTable(
                        [[Paragraph('', S_BODY), content_paras]],
                        colWidths=[accent_cell_w, content_cell_w]
                    )
                    callout_tbl.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (0, -1), HexColor(BRAND_HIGHLIGHT)),
                        ('BACKGROUND', (1, 0), (1, -1), HexColor(BRAND_LIGHT)),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('LEFTPADDING', (1, 0), (1, -1), 12),
                        ('RIGHTPADDING', (1, 0), (1, -1), 12),
                        ('TOPPADDING', (1, 0), (1, -1), 10),
                        ('BOTTOMPADDING', (1, 0), (1, -1), 10),
                        ('LEFTPADDING', (0, 0), (0, -1), 0),
                        ('RIGHTPADDING', (0, 0), (0, -1), 0),
                        ('TOPPADDING', (0, 0), (0, -1), 0),
                        ('BOTTOMPADDING', (0, 0), (0, -1), 0),
                    ]))
                    story.append(Spacer(1, 6))
                    story.append(callout_tbl)
                    story.append(Spacer(1, 6))
            else:
                tbl = build_table(tbl_obj)
                if tbl:
                    story.append(Spacer(1, 6))
                    story.append(tbl)
                    story.append(Spacer(1, 8))

    # Build
    pdf.build(story)
    print(f"OK: {input_path} → {output_path} ({os.path.getsize(output_path)//1024} KB)")


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python docx_to_pdf.py <input.docx> <output.pdf>")
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])
